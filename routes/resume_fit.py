"""Resume endpoints for HR AI Platform

Improvements:
- DOCX parsing via `python-docx`
- PDF text fallbacks: PyPDF2 -> pdfplumber -> OCR (pytesseract + OpenCV)
- Image OCR for uploaded images
- New `/resume/analyze` endpoint that returns skills and role predictions
"""
import os
from typing import List
import PyPDF2
from fastapi import APIRouter, File, UploadFile, HTTPException, Depends, Query, Form
from .. import db
from ..models import User, Resume, Job, JobMatch
from ..services import embedding_service, skill_extractor, qdrant_utils
from ..services.auth import get_current_user

import io
import logging

logger = logging.getLogger(__name__)

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    import cv2
except Exception:
    cv2 = None

try:
    import numpy as np
except Exception:
    np = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

router = APIRouter(prefix="/resume", tags=["Resume"])


def extract_text_from_path(path: str, filename: str) -> str:
    """Top-level helper for extracting text from different file types.

    Uses PyPDF2 -> pdfplumber -> OCR (pytesseract+OpenCV) for PDFs,
    python-docx for DOCX, and pytesseract+OpenCV for images.
    """
    text = ""
    ext = os.path.splitext(filename)[1].lower()

    # PDF: try PyPDF2, then pdfplumber, then OCR
    if ext == ".pdf":
        try:
            with open(path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                    except Exception:
                        continue
        except Exception:
            logger.exception("PyPDF2 failed to read PDF")

        if not text.strip() and pdfplumber is not None:
            try:
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        try:
                            p_text = page.extract_text() or ""
                            text += p_text + "\n"
                        except Exception:
                            continue
            except Exception:
                logger.exception("pdfplumber failed to read PDF")

        # If still empty, try OCR on each page. Prefer PyMuPDF (fitz) for
        # higher-quality rasterization; fall back to pdfplumber image rendering.
        if not text.strip() and pytesseract is not None and np is not None:
            DPI = int(os.getenv("PDF_OCR_DPI", "300"))

            # Try PyMuPDF first (high-quality rasterization)
            if fitz is not None:
                try:
                    doc = fitz.open(path)
                    for page in doc:
                        try:
                            mat = fitz.Matrix(DPI / 72.0, DPI / 72.0)
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            img_bytes = pix.tobytes("png")
                            # Convert bytes -> cv2 image if available
                            if cv2 is not None:
                                nparr = np.frombuffer(img_bytes, np.uint8)
                                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                                if img is None:
                                    continue
                                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                                gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                                ocr_text = pytesseract.image_to_string(gray)
                            else:
                                # PIL fallback
                                from PIL import Image
                                from io import BytesIO
                                img_pil = Image.open(BytesIO(img_bytes)).convert('RGB')
                                ocr_text = pytesseract.image_to_string(img_pil)
                            text += ocr_text + "\n"
                        except Exception:
                            continue
                except Exception:
                    logger.exception("PyMuPDF rendering + OCR failed; falling back to pdfplumber")

            # Fallback to pdfplumber rendering with configurable DPI
            if not text.strip() and pdfplumber is not None and cv2 is not None:
                try:
                    with pdfplumber.open(path) as pdf:
                        for page in pdf.pages:
                            try:
                                pil_image = page.to_image(resolution=DPI).original
                                img = np.array(pil_image.convert('RGB'))
                                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                                gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                                ocr_text = pytesseract.image_to_string(gray)
                                text += ocr_text + "\n"
                            except Exception:
                                continue
                except Exception:
                    logger.exception("pdfplumber OCR fallback failed")

    # DOCX parsing
    elif ext == ".docx":
        if DocxDocument is not None:
            try:
                doc = DocxDocument(path)
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                text = "\n".join(paragraphs)
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text += "\n" + cell.text
            except Exception:
                logger.exception("Failed to parse DOCX file")
        else:
            logger.warning("python-docx not available; cannot parse DOCX")

    # Images: OCR using pytesseract + OpenCV
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        if pytesseract is not None and cv2 is not None and np is not None:
            try:
                img = cv2.imread(path)
                if img is None:
                    # Try PIL fallback
                    from PIL import Image
                    img = np.array(Image.open(path).convert('RGB'))
                    img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                text = pytesseract.image_to_string(gray)
            except Exception:
                logger.exception("Image OCR failed")
        else:
            logger.warning("OCR dependencies missing; cannot OCR image")

    else:
        logger.warning("Unsupported file extension for text extraction: %s", ext)

    return text or ""

@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Upload and process a resume"""
    # Validate file
    allowed_types = [".pdf", ".doc", ".docx"]
    if not any(file.filename.endswith(ext) for ext in allowed_types):
        raise HTTPException(status_code=400, detail="Invalid file type")
    
    # Save file
    file_path = f"uploads/resumes/{current_user.id}_{file.filename}"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    # Extract text using top-level helper
    resume_text = extract_text_from_path(file_path, file.filename)

    # Detect language and extract skills (language-aware)
    try:
        from ..services.language_utils import detect_language
        lang = detect_language(resume_text)
    except Exception:
        lang = None

    skills = skill_extractor.extract_skills(resume_text, lang=lang)
    vector = embedding_service.embed_text(resume_text)

    # Convert vector to bytes hex safely
    try:
        import numpy as _np
        vec_np = _np.array(vector, dtype=_np.float32)
        vector_hex = vec_np.tobytes().hex()
    except Exception:
        # Fallback: JSON string
        import json
        vector_hex = json.dumps(vector)

    # Store in database
    conn = db.get_db_connection()
    cursor = conn.cursor()
    query = """
        INSERT INTO resumes (user_id, resume_text, skills, vector_embedding)
        VALUES (%s, %s, %s, %s)
    """
    cursor.execute(query, (
        current_user.id,
        resume_text,
        ",".join(skills),
        vector_hex
    ))
    resume_id = cursor.lastrowid
    conn.commit()

    cursor.close()
    db.close_db_connection(conn)

    return {
        'id': resume_id,
        'user_id': current_user.id,
        'skills': skills
    }


@router.post("/analyze")
async def analyze(
    file: UploadFile = File(None),
    resume_text: str = Form(None),
    job_description: str = Form(None),
    current_user: User = Depends(get_current_user)
):
    """Analyze resume (file upload or raw text) and return skills + role prediction.

    - Accepts either an uploaded file or raw `resume_text` form field.
    - Optionally accepts `job_description` to compute match score.
    """
    if not file and not resume_text:
        raise HTTPException(status_code=400, detail="Provide a resume file or resume_text")

    if file:
        # Save and extract text
        file_path = f"uploads/resumes/{current_user.id}_{file.filename}"
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())

        resume_text = extract_text_from_path(file_path, file.filename)

    # Detect language and extract skills (language-aware)
    try:
        from ..services.language_utils import detect_language
        lang = detect_language(resume_text)
    except Exception:
        lang = None

    skills = skill_extractor.extract_skills(resume_text, lang=lang)

    # Role prediction (baseline TF-IDF + LogisticRegression)
    role_scores = []
    try:
        from ..ai_engine.role_predictor import role_predictor
        role_scores = role_predictor.predict(resume_text, top_k=3)
    except Exception:
        logger.exception("Role predictor unavailable or failed")

    result = {
        'skills': skills,
        'role_predictions': role_scores,
        'resume_text_preview': (resume_text[:200] + '...') if resume_text and len(resume_text) > 200 else (resume_text or "")
    }

    if lang:
        result['language'] = lang

    if job_description:
        emb_resume = embedding_service.embed_text(resume_text)
        emb_job = embedding_service.embed_text(job_description)
        match_score = embedding_service.calculate_similarity(emb_resume, emb_job)
        result['match_score'] = float(match_score)

    return result


@router.post('/upload-voice')
async def upload_voice(
    file: UploadFile = File(...),
    language: str | None = Form(None),
    current_user: User = Depends(get_current_user)
):
    """Upload an audio file (wav/mp3) and transcribe it to resume text, then analyze."""
    # Save audio file
    allowed = {'.wav', '.mp3', '.m4a', '.ogg', '.flac', '.webm', '.aac'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(status_code=400, detail='Unsupported audio format')
    file_path = f"uploads/resumes/{current_user.id}_audio_{file.filename}"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, 'wb') as f:
        f.write(await file.read())

    # Transcribe
    try:
        # local import to avoid circular imports
        from ..services.speech_to_text import transcribe_audio
        resume_text = transcribe_audio(file_path, language=language)
    except Exception as e:
        logger.exception('Speech-to-text failed')
        raise HTTPException(status_code=500, detail=str(e))

    # Reuse analyze logic: extract skills and role predictions
    try:
        from ..services.language_utils import detect_language
        detected = detect_language(resume_text)
    except Exception:
        detected = None

    skills = skill_extractor.extract_skills(resume_text, lang=detected)
    role_scores = []
    try:
        from ..ai_engine.role_predictor import role_predictor
        role_scores = role_predictor.predict(resume_text, top_k=3)
    except Exception:
        logger.exception('Role predictor unavailable')

    result = {
        'skills': skills,
        'role_predictions': role_scores,
        'resume_text': resume_text
    }
    if detected:
        result['language'] = detected
    return result


@router.post('/chat')
async def chat_resume(
    resume_id: int = Form(...),
    message: str = Form(...),
    current_user: User = Depends(get_current_user)
):
    """Simple chatbot against a stored resume. Returns extracted answer.

    If an LLM is configured the endpoint will try to use it for a richer
    response; otherwise a retrieval-based answer is returned.
    """
    # Fetch resume text from DB
    conn = db.get_db_connection()
    cursor = conn.cursor()
    query = "SELECT resume_text FROM resumes WHERE id = %s AND user_id = %s"
    cursor.execute(query, (resume_id, current_user.id))
    row = cursor.fetchone()
    cursor.close()
    db.close_db_connection(conn)
    if not row:
        raise HTTPException(status_code=404, detail='Resume not found')
    resume_text = row[0] or ""

    # Try LLM-assisted answer
    try:
        from services.llm import call_llm
        prompt = f"You are an assistant that answers questions about a candidate's resume.\nResume:\n{resume_text}\n\nQuestion: {message}\nAnswer concisely."
        ans = call_llm(prompt, max_tokens=256)
        return {'answer': ans}
    except NotImplementedError:
        # Gemini stub not implemented; fall through to retrieval
        pass
    except Exception:
        logger.debug('LLM call failed; falling back to retrieval', exc_info=True)

    # Retrieval-based fallback
    try:
        from services.qa_service import answer_from_resume
        ans = answer_from_resume(resume_text, message)
        return {'answer': ans}
    except Exception:
        logger.exception('QA failed')
        raise HTTPException(status_code=500, detail='Failed to answer')


@router.post('/rank')
async def rank_resumes(
    job_description: str = Form(...),
    top_k: int = Form(10),
    current_user: User = Depends(get_current_user)
):
    """Rank all resumes in DB by relevance to the given job description."""
    conn = db.get_db_connection()
    cursor = conn.cursor()
    query = "SELECT id, user_id, resume_text, skills, vector_embedding FROM resumes"
    cursor.execute(query)
    rows = cursor.fetchall()

    # Compute embedding for job
    emb_job = embedding_service.embed_text(job_description)

    results = []
    for r in rows:
        rid = r[0]
        resume_text = r[2] or ''
        # Compute similarity
        try:
            emb_resume = embedding_service.embed_text(resume_text)
            score = embedding_service.calculate_similarity(emb_resume, emb_job)
        except Exception:
            score = 0.0
        results.append({'resume_id': rid, 'user_id': r[1], 'score': float(score), 'skills': (r[3] or '').split(',')})

    # Sort and return top_k
    results.sort(key=lambda x: x['score'], reverse=True)
    return {'ranked': results[:top_k]}


@router.post('/suggest')
async def suggest_improvements(
    resume_id: int | None = Form(None),
    resume_text: str | None = Form(None),
    job_description: str | None = Form(None),
    current_user: User = Depends(get_current_user)
):
    """Return automated improvement suggestions for a resume.

    Accepts either `resume_id` (existing resume) or raw `resume_text`.
    Optionally accepts `job_description` to tailor suggestions.
    """
    text = resume_text
    if resume_id and not text:
        # fetch from DB
        conn = db.get_db_connection()
        cursor = conn.cursor()
        query = "SELECT resume_text FROM resumes WHERE id = %s AND user_id = %s"
        cursor.execute(query, (resume_id, current_user.id))
        row = cursor.fetchone()
        cursor.close()
        db.close_db_connection(conn)
        if not row:
            raise HTTPException(status_code=404, detail='Resume not found')
        text = row[0] or ''

    if not text:
        raise HTTPException(status_code=400, detail='Provide resume_id or resume_text')

    try:
        from services.resume_advisor import suggest_improvements
        suggestions = suggest_improvements(text, job_description)
        return {'suggestions': suggestions}
    except Exception:
        logger.exception('Failed to generate suggestions')
        raise HTTPException(status_code=500, detail='Failed to generate suggestions')

@router.get("/my-resumes", response_model=List[dict])
async def get_my_resumes(current_user: User = Depends(get_current_user)):
    """Get all resumes for current user"""
    conn = db.get_db_connection()
    cursor = conn.cursor()
    query = "SELECT * FROM resumes WHERE user_id = %s"
    cursor.execute(query, (current_user.id,))
    resumes = cursor.fetchall()
    cursor.close()
    db.close_db_connection(conn)
    
    return [{
        'id': resume[0],
        'user_id': resume[1],
        'skills': resume[3].split(",") if resume[3] else []
    } for resume in resumes]

@router.post("/{resume_id}/match")
async def match_resume(
    resume_id: int,
    top_k: int = Query(default=3, gt=0),
    current_user: User = Depends(get_current_user)
):
    """Match a resume with available jobs"""
    conn = db.get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    query = "SELECT * FROM resumes WHERE id = %s AND user_id = %s"
    cursor.execute(query, (resume_id, current_user.id))
    resume = cursor.fetchone()
    
    if not resume:
        cursor.close()
        db.close_db_connection(conn)
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get all jobs
    query = "SELECT * FROM jobs"
    cursor.execute(query)
    jobs = cursor.fetchall()
    
    matches = []
    for job in jobs:
        # Calculate similarity score
        similarity = embedding_service.calculate_similarity(
            resume[4],  # vector_embedding
            job[3]     # description
        )
        
        # Calculate skill overlap
        resume_skills = set(resume[3].split(","))
        job_skills = set(job[5].split(","))  # required_skills
        skill_overlap = len(resume_skills.intersection(job_skills))
        skill_overlap_score = skill_overlap / len(job_skills) if job_skills else 0
        
        # Store match in database
        query = """
            INSERT INTO job_matches (job_id, resume_id, match_score, skills_matched)
            VALUES (%s, %s, %s, %s)
        """
        cursor.execute(query, (
            job[0],  # job_id
            resume_id,
            similarity,
            ",".join(resume_skills.intersection(job_skills))
        ))
        
        matches.append({
            "job": {
                "id": job[0],
                "title": job[1],
                "description": job[3],
                "required_skills": list(job_skills)
            },
            "match_score": float(similarity),
            "skill_overlap_score": float(skill_overlap_score),
            "matching_skills": list(resume_skills.intersection(job_skills)),
            "missing_skills": list(job_skills - resume_skills)
        })
    
    conn.commit()
    cursor.close()
    db.close_db_connection(conn)
    
    # Sort matches by score
    matches.sort(key=lambda x: x["match_score"], reverse=True)
    
    return matches[:top_k]

@router.delete("/{resume_id}")
async def delete_resume(resume_id: int, current_user: User = Depends(get_current_user)):
    """Delete a resume"""
    conn = db.get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    query = "SELECT * FROM resumes WHERE id = %s AND user_id = %s"
    cursor.execute(query, (resume_id, current_user.id))
    resume = cursor.fetchone()
    
    if not resume:
        cursor.close()
        db.close_db_connection(conn)
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Delete from .database
    query = "DELETE FROM resumes WHERE id = %s"
    cursor.execute(query, (resume_id,))
    conn.commit()
    cursor.close()
    db.close_db_connection(conn)
    
    # Delete associated file
    file_path = f"uploads/resumes/{current_user.id}_{resume_id}"
    if os.path.exists(file_path):
        os.remove(file_path)
    
    return {"message": "Resume deleted successfully"}
